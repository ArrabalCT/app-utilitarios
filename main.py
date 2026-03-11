import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import io
import os
import re
from PIL import Image, ImageOps

# --- Funções de Formatação do Word ---
def adicionar_borda_inferior(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def adicionar_campo_numpages(paragraph):
    p = paragraph._p
    r1 = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r1.append(fldChar1)
    p.append(r1)
    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' NUMPAGES '
    r2.append(instrText)
    p.append(r2)
    r3 = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    r3.append(fldChar2)
    p.append(r3)
    r4 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = 'xx'
    r4.append(t)
    p.append(r4)
    r5 = OxmlElement('w:r')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r5.append(fldChar3)
    p.append(r5)

# --- Configuração da Página ---
st.set_page_config(page_title="Vistoria - GTA", page_icon="📋", layout="centered")

# --- Variáveis de Sessão ---
if 'fotos' not in st.session_state: st.session_state['fotos'] = []
if 'camera_key' not in st.session_state: st.session_state['camera_key'] = 0 
if 'mk' not in st.session_state: st.session_state['mk'] = 0 

mk = st.session_state['mk']

# --- Listas de Dados ---
delegados = ["Adilson Antonio Marcondes dos Santos", "Adriane Goncalves", "Anisio Galdioli", "Cesar Aparecido Vieira da Silva", "Daniel Souza Baptista de Castro", "Ernani Ronaldo Giannico Braga", "Fabio Germano Figueiredo Cabett", "Flavia Maria Rocha Rollo", "Francisco Sannini Neto", "Hugo Parreiras de Macedo", "Jose Marcelo Silva Hial", "Leonardo da Costa Ferreira", "Marcelo Vieira Cavalcante", "Mario Celso Ribeiro Senne", "Paulo Roberto Gruschka Castilho", "Paulo Sergio Barbosa", "Pedro Rossati", "Sergio Lucas Adler Guedes de Oliveira", "Vania Idalira Z. de Oliveira", "Outro..."]
peritos = ["Alexandre Rabello de Oliveira", "Bruna Fernandes Nogueira", "Claude Thiago Arrabal", "Jéssica Pereira Gonçalves", "Júlia Soares Melo", "Luiz Fausto Prado Vasques", "Marcelo Mourão Dantas", "Márcio Steinmetz Soares", "Sarah Costa Teixeira", "Ruan Carvalho de Souza"]
cidades = ["Aparecida", "Cachoeira Paulista", "Canas", "Cunha", "Guaratinguetá", "Lorena", "Piquete", "Potim", "Roseira", "Outra..."]
dps_por_cidade = {"Aparecida": ["DEL.POL.APARECIDA"], "Canas": ["DEL.POL.CANAS"], "Cachoeira Paulista": ["DEL.POL.CACHOEIRA PAULISTA"], "Cunha": ["DEL.POL.CUNHA"], "Guaratinguetá": ["01º D.P. GUARATINGUETA", "02º D.P. GUARATINGUETA", "DEL.SEC.GUARATINGUETA PLANTÃO", "DISE- DEL.SEC.GUARATINGUETA"], "Lorena": ["01º D.P. LORENA", "02º D.P. LORENA", "DEL.POL.LORENA"], "Piquete": ["DEL.POL.PIQUETE"], "Potim": ["DEL.POL.POTIM"], "Roseira": ["DEL.POL.ROSEIRA"]}

# --- INTERFACE ---
st.title("Gerador de Laudos - Vistoria")

st.header("1. Cabeçalho e Identificação")
colBO1, colBO2 = st.columns(2)
with colBO1: 
    bo_input = st.text_input("Número do BO:", value="", placeholder="Ex: LT0644", key=f"bo_{mk}").upper()

with colBO2: bo_ano = st.text_input("Ano do BO:", value="2026", max_chars=4, key=f"ano_{mk}")

data_selecionada = st.date_input("Data do Laudo:", format="DD/MM/YYYY", key=f"data_{mk}")
meses = {1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril', 5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto', 9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'}
data_extenso = f"{data_selecionada.day} de {meses[data_selecionada.month]} de {data_selecionada.year}"

objetivos_selecionados = st.multiselect("Objetivo da Perícia:", ["Vistoria", "Fotografação", "Constatação de Danos", "Verificação dos Sistemas de Segurança", "Adulteração de Sinais Identificadores"], default=["Vistoria", "Constatação de Danos"], key=f"obj_{mk}")
perito_selecionado = st.selectbox("Perito Criminal:", peritos, index=peritos.index("Claude Thiago Arrabal"), key=f"per_{mk}")

del_sel = st.selectbox("Autoridade Policial:", delegados, index=delegados.index("Adilson Antonio Marcondes dos Santos") if "Adilson Antonio Marcondes dos Santos" in delegados else 0, key=f"del_sel_{mk}")
if del_sel == "Outro...":
    delegado_selecionado = st.text_input("Digite o nome da Autoridade Policial:", key=f"del_dig_{mk}")
else:
    delegado_selecionado = del_sel

colC1, colC2 = st.columns(2)
with colC1: 
    cid_sel = st.selectbox("Cidade:", cidades, index=cidades.index("Guaratinguetá") if "Guaratinguetá" in cidades else 0, key=f"cid_sel_{mk}")
    if cid_sel == "Outra...":
        cidade_selecionada = st.text_input("Digite o nome da Cidade:", key=f"cid_dig_{mk}")
    else:
        cidade_selecionada = cid_sel

with colC2:
    if cid_sel == "Outra...":
        delegacia_selecionada = st.text_input("Digite o nome da Delegacia:", key=f"dp_dig_{mk}")
    else:
        opcoes_dp = dps_por_cidade[cid_sel] + ["Outra..."]
        index_padrao = opcoes_dp.index("DEL.SEC.GUARATINGUETA PLANTÃO") if "DEL.SEC.GUARATINGUETA PLANTÃO" in opcoes_dp else 0
        dp_sel = st.selectbox("Delegacia:", opcoes_dp, index=index_padrao, key=f"dp_sel_{mk}")
        if dp_sel == "Outra...":
            delegacia_selecionada = st.text_input("Digite o nome da Delegacia:", key=f"dp_dig_esp_{mk}")
        else:
            delegacia_selecionada = dp_sel

st.header("2. Veículo")
colT1, colT2, colT3 = st.columns(3)
with colT1: tipo_v = st.selectbox("Tipo:", ["Automóvel", "Motocicleta", "Caminhonete", "Caminhoneta", "Caminhão", "Ônibus", "Micro-ônibus", "Reboque", "Semirreboque", "Ciclomotor", "Outro"], key=f"tv_{mk}")
with colT2: esp_v = st.selectbox("Espécie:", ["Passageiro", "Carga", "Tração"], key=f"ev_{mk}")
with colT3: cat_v = st.selectbox("Categoria:", ["Particular", "Aluguel", "Oficial", "Aprendizagem"], key=f"cv_{mk}")

colV1, colV2, colV3 = st.columns(3)
with colV1: placa = st.text_input("Placa:", value="", placeholder="Ex: ABC1D23", key=f"placa_{mk}").upper()
with colV3: mod_v = st.text_input("Marca/Modelo:", value="", placeholder="Ex: FORD FOCUS SE", key=f"mod_{mk}").upper()
with colV2:
    cores_basicas = ["Amarela", "Azul", "Branca", "Cinza", "Marrom", "Prata", "Preta", "Verde", "Vermelha", "Outra..."]
    cor_sel = st.selectbox("Cor:", cores_basicas, index=cores_basicas.index("Prata"), key=f"csel_{mk}")
    cor = st.text_input("Qual cor?", key=f"cdig_{mk}") if cor_sel == "Outra..." else cor_sel

st.header("3. Constatações")
# Listas padrão de peças organizadas por região
pecas_padrao = {
    "Dianteira": ["Para-choque", "Grade", "Emblema", "Capô", "Farol Esquerdo", "Farol Direito", "Para-lama Esquerdo", "Para-lama Direito", "Para-brisa", "Placa Dianteira"],
    "Traseira": ["Para-choque", "Tampa do Porta-malas", "Lanterna Esquerda", "Lanterna Direita", "Vidro Traseiro", "Placa Traseira"],
    "Lateral Esquerda": ["Porta Dianteira", "Porta Traseira", "Retrovisor", "Soleira", "Vidros", "Colunas", "Roda/Pneu"],
    "Lateral Direita": ["Porta Dianteira", "Porta Traseira", "Retrovisor", "Soleira", "Vidros", "Colunas", "Roda/Pneu"],
    "Teto": ["Teto", "Rack/Longarina"]
}

# --- 3. Constatações (CORRIGIDO PARA EVITAR CHAVES DUPLICADAS) ---
regioes_detalhes = {}
for regiao, lista in pecas_padrao.items():
    st.markdown(f"##### 🚗 {regiao}")
    selecionadas = st.multiselect(f"Peças na {regiao}:", lista, key=f"sel_{regiao}_{mk}")
    outras = st.text_input(f"Outras na {regiao}:", placeholder="Ex: Friso, Cárter...", key=f"out_{regiao}_{mk}")
    
    todas = selecionadas + [p.strip() for p in outras.split(",") if p.strip()]
    regioes_detalhes[regiao] = {}
    
    for peca in todas:
        st.markdown(f"**Detalhes: {peca}**")
        c1, c2, c3 = st.columns(3)
        with c1: 
            t = st.multiselect("Dano:", ["Amolgamento", "Fratura", "Atritamento", "Quebra"], 
                               key=f"t_{peca}_{regiao}_{mk}")
        with c2: 
            o = st.multiselect("Sentido:", ["Da frente para trás", "De trás para a frente", "Da esquerda para a direita", "Da direita para a esquerda", "De cima para baixo", "De baixo para cima", "De fora para dentro", "De dentro para fora"], 
                               key=f"o_{peca}_{regiao}_{mk}")
        with c3: 
            a = st.multiselect("Altura:", ["Terço superior", "Terço médio", "Terço inferior"], 
                               key=f"a_{peca}_{regiao}_{mk}")
        regioes_detalhes[regiao][peca] = {"tipo": t, "ori": o, "alt": a}
    
st.markdown("##### Componentes Adicionais")
colE1, colE2 = st.columns(2)
with colE1: 
    s_ele = st.selectbox("Sistemas Elétricos:", ["Funcionando a contento", "Avariados devido ao impacto", "Sem bateria", "Ausência de chaves"], key=f"se_{mk}")
with colE2: 
    s_fre = st.selectbox("Sistema de Freios:", ["Operantes em exame estático", "Inoperantes", "Inoperantes face aos danos", "Travados em decorrência do impacto", "Prejudicado"], key=f"sf_{mk}")

pneus = st.multiselect("Pneus com desgaste acentuado (Lisos / TWI atingido):", ["Dianteiro Esquerdo", "Dianteiro Direito", "Traseiro Esquerdo", "Traseiro Direito", "Estepe"], key=f"pn_{mk}")
consideracoes = st.text_input("Outras Considerações:", placeholder="Ex: Veículo recolhido por guincho...", key=f"cons_{mk}")

# --- GERADOR DE TEXTO ---
def montar_regiao(nome, dados):
    if not dados: return ""
    linhas = []
    for peca, det in dados.items():
        t = ", ".join(det['tipo']).lower() if det['tipo'] else "avaria(s)"
        a = f" (altura: {', '.join(det['alt']).lower()})" if det['alt'] else ""
        o = f" (sentido: {', '.join(det['ori']).lower()})" if det['ori'] else ""
        linhas.append(f"{peca} com {t}{a}{o}")
    if not linhas: return ""
    return f"• {nome}: {'; '.join(linhas[:-1]) + ' e ' + linhas[-1] if len(linhas) > 1 else linhas[0]}.\n"

# Forçando Maiúsculas para as variáveis principais do texto
cat_str = cat_v.split(" ")[0].lower()
marca_modelo_str = mod_v.upper() if mod_v else "[MARCA/MODELO]"
placa_str = placa.upper() if placa else "[PLACA]"
cor_str = cor.lower() if cor else "[COR]"

# Texto contínuo: sem \n\n antes de "Nas inspeções..."
txt_gerado = f"Trata-se de um veículo do tipo {tipo_v.lower()}, espécie {esp_v.lower()}, categoria {cat_str}, marca/modelo {marca_modelo_str}, de cor {cor_str}, ostentando a placa {placa_str}. "

if any(regioes_detalhes.values()):
    txt_gerado += "Nas inspeções realizadas, constatou-se que a unidade apresentava avarias recentes localizadas nas seguintes regiões/peças:\n"
    for r, d in regioes_detalhes.items(): txt_gerado += montar_regiao(r, d)

# Adicionando os Sistemas na mesma linha (separados por espaço em vez de \n)
txt_gerado += f"Quanto aos sistemas elétricos, encontravam-se {s_ele.lower()}. "
txt_gerado += f"O sistema de freios apresentou-se {s_fre.lower()}. "

if not pneus:
    txt_gerado += "Os pneumáticos encontravam-se em aparente bom estado de conservação."
elif len(pneus) >= 4 and "Estepe" not in pneus:
    txt_gerado += "Todos os pneumáticos de rodagem apresentavam desgaste acentuado da banda de rodagem, atingindo os indicadores de desgaste (TWI)."
else:
    pneus_str = ", ".join(pneus).lower()
    if "," in pneus_str: pneus_str = " e ".join(pneus_str.rsplit(", ", 1))
    plural = "apresentavam" if len(pneus) > 1 else "apresentava"
    txt_gerado += f"Os pneumáticos encontravam-se em aparente bom estado de conservação, com exceção do(s) {pneus_str}, que {plural} desgaste acentuado da banda de rodagem, atingindo ou ultrapassando os indicadores de desgaste (TWI)."

if consideracoes:
    txt_gerado += f"\n\nConsiderações Adicionais: {consideracoes}"

# Sincronização da Caixa Editável
if st.session_state.get(f"track_{mk}") != txt_gerado:
    st.session_state[f"edit_{mk}"] = txt_gerado
    st.session_state[f"track_{mk}"] = txt_gerado

st.header("4. Edição e Word")
st.warning("⚠️ Pode digitar diretamente na caixa abaixo. No entanto, deixe as edições manuais para o **FINAL**. Se alterar as opções em cima, o sistema irá recriar a frase e apagar as suas edições!")
texto_final = st.text_area("Texto final que vai para o Laudo:", height=300, key=f"edit_{mk}")

# Fotos
foto = st.camera_input("Tirar fotografia", key=f"cam_{st.session_state['camera_key']}")
if foto:
    colA, colR = st.columns(2)
    with colA:
        if st.button("✅ ACEITAR FOTO", type="primary", use_container_width=True):
            img = Image.open(io.BytesIO(foto.getvalue()))
            st.session_state['fotos'].append({'img': ImageOps.exif_transpose(img)})
            st.session_state['camera_key'] += 1
            st.rerun()
    with colR:
        if st.button("❌ REJEITAR", use_container_width=True):
            st.session_state['camera_key'] += 1
            st.rerun()

fotos_up = st.file_uploader("Ou carregue da galeria", type=['jpg', 'jpeg', 'png'], accept_multiple_files=True, key=f"up_{mk}")
if fotos_up:
    for f in fotos_up:
        img = Image.open(io.BytesIO(f.getvalue()))
        st.session_state['fotos'].append({'img': ImageOps.exif_transpose(img)})
    st.rerun()

if st.session_state['fotos']:
    st.markdown("### 📸 Fotos Anexadas")
    cols = st.columns(3)
    for i, foto_data in enumerate(st.session_state['fotos']):
        with cols[i % 3]:
            st.image(foto_data['img'], use_column_width=True)
            if st.button("❌ Apagar", key=f"del_{i}_{mk}", type="secondary"):
                st.session_state['fotos'].pop(i)
                st.rerun()

# --- BOTÕES FINAIS ---
st.header("5. Finalizar")
c1, c2 = st.columns(2)

with c1:
    if st.button("Criar Laudo (.docx)", type="primary", use_container_width=True):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Courier New'
        style.font.size = Pt(11)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Cabeçalho
        section = doc.sections[0]
        header = section.header
        for p in header.paragraphs: p.text = ""
        
        # Ajustando a largura total da tabela para 15.5 cm para respeitar as margens
        table = header.add_table(rows=1, cols=3, width=Cm(15.5))
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        largura_lateral = Cm(2.2)
        largura_meio = Cm(11.1)

        table.columns[0].width = largura_lateral
        table.columns[1].width = largura_meio
        table.columns[2].width = largura_lateral
        
        for cell in table.columns[0].cells: cell.width = largura_lateral
        for cell in table.columns[1].cells: cell.width = largura_meio
        for cell in table.columns[2].cells: cell.width = largura_lateral

        p_left = table.cell(0, 0).paragraphs[0]; p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if os.path.exists("logo_ssp.png"): p_left.add_run().add_picture("logo_ssp.png", width=Cm(1.8))
        
        p_center = table.cell(0, 1).paragraphs[0]; p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_h1 = p_center.add_run("SECRETARIA DA SEGURANÇA PÚBLICA\nSUPERINTENDÊNCIA DA POLÍCIA TÉCNICO-CIENTÍFICA\n")
        run_h1.font.size = Pt(11); run_h1.bold = False
        run_h2 = p_center.add_run("INSTITUTO DE CRIMINALÍSTICA\n“PERITO CRIMINAL DR. OCTÁVIO EDUARDO DE BRITO ALVARENGA”\nNÚCLEO DE PERÍCIAS CRIMINALÍSTICAS DE SÃO JOSÉ DOS CAMPOS\nEQUIPE DE PERÍCIAS CRIMINALÍSTICAS DE GUARATINGUETÁ")
        run_h2.font.size = Pt(8); run_h2.bold = False

        p_right = table.cell(0, 2).paragraphs[0]; p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if os.path.exists("logo_ic.png"): p_right.add_run().add_picture("logo_ic.png", width=Cm(1.8))

        if bo_input:
            p_bo = doc.add_paragraph()
            p_bo.add_run(f"BO {bo_input.upper()} / {bo_ano} - {delegacia_selecionada}")
            p_bo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Corpo
        p_nat = doc.add_paragraph()
        run = p_nat.add_run("1 – NATUREZA: Vistoria"); run.bold = True; run.font.size = Pt(14)
        adicionar_borda_inferior(p_nat)
        
        preambulo = (f"Aos {data_extenso}, no Instituto de Criminalística, da Superintendência da Polícia Técnico-Científica, "
                     f"da Secretaria da Segurança Pública do Estado de São Paulo, de conformidade com o disposto no artigo 178 "
                     f"do Decreto-Lei nº. 3689, de 03 de outubro de 1941, pelo Diretor do Instituto de Criminalística, Ricardo Lopes Ortega, "
                     f"foi designado o Perito Criminal {perito_selecionado}, para proceder ao exame supracitado, em atendimento à requisição "
                     f"da Autoridade Policial, Dr(a). {delegado_selecionado}, titular/em exercício na {delegacia_selecionada}.")
        doc.add_paragraph(preambulo)

        p_obj = doc.add_paragraph()
        run_obj = p_obj.add_run("2 - OBJETIVO DA PERÍCIA:"); run_obj.bold = True; run_obj.font.size = Pt(14)
        adicionar_borda_inferior(p_obj)
        objetivos_str = ", ".join(objetivos_selecionados) if objetivos_selecionados else "Não especificado"
        doc.add_paragraph(f"Consta na requisição de exame: {objetivos_str}.")

        p_ex = doc.add_paragraph()
        run2 = p_ex.add_run("3 – DOS EXAMES:"); run2.bold = True; run2.font.size = Pt(14)
        adicionar_borda_inferior(p_ex)
        
        for linha in texto_final.split("\n"): doc.add_paragraph(linha.strip())
            
        if st.session_state['fotos']:
            doc.add_page_break()
            p_fotos = doc.add_paragraph()
            run_fotos = p_fotos.add_run("4 - REGISTRO FOTOGRÁFICO"); run_fotos.bold = True; run_fotos.font.size = Pt(14)
            adicionar_borda_inferior(p_fotos)
            
            for i, foto_data in enumerate(st.session_state['fotos']):
                img = foto_data['img']
                if img.mode != 'RGB': img = img.convert('RGB')
                largura_foto = Cm(14.0) if img.width > img.height else Cm(9.5)
                buf = io.BytesIO()
                img.save(buf, format='JPEG', quality=90)
                buf.seek(0)
                
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(buf, width=largura_foto)
                
                legenda = doc.add_paragraph(f"Fotografia {i+1}: Vistoria do veículo.")
                legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("")
            
        # Encerramento
        p_relatar = doc.add_paragraph("Era o que havia a relatar.")
        p_relatar.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_paginas = doc.add_paragraph("Este laudo vai impresso em ")
        adicionar_campo_numpages(p_paginas)  
        p_paginas.add_run(" páginas, além da capa, ficando arquivada cópia digital no sistema GDL da SPTC.")
        p_paginas.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        p_assinatura = doc.add_paragraph(); p_assinatura.paragraph_format.space_after = p_assinatura.paragraph_format.space_before = Pt(0)
        p_assinatura.add_run(perito_selecionado).bold = True
        p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_cargo = doc.add_paragraph("Perito Criminal Relator"); p_cargo.paragraph_format.space_after = p_cargo.paragraph_format.space_before = Pt(0)
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buf_doc = io.BytesIO(); doc.save(buf_doc); buf_doc.seek(0)
        nome_arquivo = f"Laudo_BO_{bo_input}_{bo_ano}.docx" if bo_input else "Laudo_Sem_BO.docx"
        st.download_button("⬇️ Descarregar Laudo Final", buf_doc, nome_arquivo, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

with c2:
    if st.button("🔄 Novo Veículo (Limpar Tudo)", type="secondary", use_container_width=True):
        current_mk = st.session_state.get('mk', 0)
        st.session_state.clear()
        st.session_state['mk'] = current_mk + 1
        st.rerun()
